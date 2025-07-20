"""
Codebase to Text Converter

This module provides functionality to convert codebases (folder structures with files)
into a single text file or Microsoft Word document (.docx), while preserving folder
structure and file contents. It supports advanced exclusion patterns and GitHub repositories.
"""

import os
import argparse
import shutil
import fnmatch
import tempfile
import sys

import git
from docx import Document
from docx.shared import Inches


class CodebaseToText:
    """
    Convert codebase to text with advanced exclusion patterns.

    This class provides functionality to convert local directories or GitHub repositories
    into text or DOCX files while supporting sophisticated exclusion patterns including
    wildcards, directory patterns, and .exclude file support.
    """

    def __init__(self, input_path, output_path, output_type, verbose=False,
                 exclude_hidden=False, exclude=None):
        """
        Initialize CodebaseToText converter.
        
        Args:
            input_path: Path to local directory or GitHub repository URL
            output_path: Path for the output file
            output_type: Output format ('txt' or 'docx')
            verbose: Enable detailed logging (default: False)
            exclude_hidden: Exclude hidden files and directories (default: False)
            exclude: List of exclusion patterns (default: None)
        """
        self.input_path = input_path
        self.output_path = output_path
        self.output_type = output_type
        self.config = {
            'verbose': verbose,
            'exclude_hidden': exclude_hidden,
        }
        self.temp_folder_path = None

        # Initialize exclusion patterns
        self.exclude_patterns = set()
        self.excluded_files_count = 0

        # Add supported image extensions
        self.image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}

        # Load exclusion patterns from various sources
        self._load_exclusion_patterns(exclude)

    @property
    def verbose(self):
        """Get verbose setting"""
        return self.config['verbose']

    @property
    def exclude_hidden(self):
        """Get exclude_hidden setting"""
        return self.config['exclude_hidden']

    def _load_exclusion_patterns(self, exclude_args):
        """Load exclusion patterns from CLI args, defaults and .exclude file."""
        self._add_cli_patterns(exclude_args)
        self._add_default_patterns()
        self._add_file_patterns()

        if self.verbose and self.exclude_patterns:
            print(f"Active exclusion patterns: {sorted(self.exclude_patterns)}")

    def _add_cli_patterns(self, exclude_args):
        """Add patterns provided via command line."""
        if not exclude_args:
            return
        for pattern in exclude_args:
            for p in pattern.split(','):
                p = p.strip()
                if p:
                    self.exclude_patterns.add(p)

    def _add_default_patterns(self):
        """Add default exclusion patterns for common files/folders."""
        default_excludes = {
            '.git/', '.git/**',
            '__pycache__/', '**/__pycache__/**',
            '*.pyc', '*.pyo', '*.pyd',
            '.venv/', 'venv/', 'env/',
            'node_modules/',
            '.DS_Store',
            '*.log', '*.tmp',
            '.pytest_cache/',
            '.coverage',
            'build/', 'dist/',
            '*.egg-info/',
        }
        self.exclude_patterns.update(default_excludes)
    def _add_file_patterns(self):
        """Load patterns from a .exclude file if present."""
        exclude_file_path = os.path.join(
            self.input_path if not self.is_github_repo() else '.',
            '.exclude'
        )
        if not os.path.exists(exclude_file_path):
            return
        try:
            with open(exclude_file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#'):
                        self.exclude_patterns.add(line)
            if self.verbose:
                print(f"Loaded exclusion patterns from {exclude_file_path}")
        except (OSError, UnicodeDecodeError) as e:
            if self.verbose:
                print(f"Warning: Could not read .exclude file: {e}")

    def _normalize_path(self, file_path, base_path):
        """Normalize path for pattern matching"""
        try:
            # Get relative path from base
            rel_path = os.path.relpath(file_path, base_path)
            # Convert to forward slashes for consistent pattern matching
            return rel_path.replace(os.sep, '/')
        except ValueError:
            # If relative path can't be computed, use absolute path
            return file_path.replace(os.sep, '/')

    def _should_exclude(self, file_path, base_path):
        """Check if file/directory should be excluded based on patterns"""
        if not self.exclude_patterns:
            return False

        # Handle hidden files if exclude_hidden is True
        if self.exclude_hidden and self._is_hidden_file(file_path):
            return True

        # Normalize the path for pattern matching
        normalized_path = self._normalize_path(file_path, base_path)
        filename = os.path.basename(file_path)

        # Check against all exclusion patterns
        for pattern in self.exclude_patterns:
            if self._pattern_matches(pattern, normalized_path, filename):
                return True

        return False

    def _pattern_matches(self, pattern, normalized_path, filename):
        """Return True if the pattern matches the given file."""
        if fnmatch.fnmatch(filename, pattern):
            return True
        if fnmatch.fnmatch(normalized_path, pattern):
            return True
        if pattern.endswith('/'):
            if self._dir_pattern_match(pattern.rstrip('/'), normalized_path):
                return True
        if '**' in pattern:
            if self._recursive_pattern_match(pattern, normalized_path):
                return True

        return False

    def _dir_pattern_match(self, dir_pattern, normalized_path):
        """Check directory style pattern."""
        for part in normalized_path.split('/'):
            if fnmatch.fnmatch(part, dir_pattern):
                return True
        return False

    def _recursive_pattern_match(self, pattern, normalized_path):
        """Handle patterns that include ** for recursion."""
        recursive_pattern = pattern.replace('**/', '').replace('**', '*')
        if fnmatch.fnmatch(normalized_path, recursive_pattern):
            return True
        path_parts = normalized_path.split('/')
        for i in range(len(path_parts)):
            partial_path = '/'.join(path_parts[i:])
            if fnmatch.fnmatch(partial_path, recursive_pattern):
                return True
        return False

    def _parse_folder(self, folder_path):
        """Parse folder structure, respecting exclusion patterns"""
        tree = ""
        excluded_dirs = set()

        for root, dirs, files in os.walk(folder_path):
            if self._handle_excluded_directory(root, folder_path, excluded_dirs):
                continue

            if self._skip_excluded_dir(root, excluded_dirs):
                continue

            self._filter_excluded_directories(dirs, root, folder_path)
            tree += self._generate_directory_entry(root, folder_path)
            tree += self._generate_file_entries(files, root, folder_path)

        self._log_tree_results(tree)
        return tree

    def _handle_excluded_directory(self, root, folder_path, excluded_dirs):
        """Handle directory exclusion and return True if directory should be skipped"""
        if self._should_exclude(root, folder_path):
            if self.verbose:
                print(f"Excluding directory: {root}")
            self.excluded_files_count += 1
            excluded_dirs.add(root)
            return True
        return False

    def _filter_excluded_directories(self, dirs, root, folder_path):
        """Filter out excluded directories from the dirs list"""
        original_dirs = dirs[:]
        dirs[:] = []
        for d in original_dirs:
            dir_path = os.path.join(root, d)
            if not self._should_exclude(dir_path, folder_path):
                dirs.append(d)
            elif self.verbose:
                print(f"Excluding directory from tree: {dir_path}")
                self.excluded_files_count += 1

        # Apply hidden file exclusion
        if self.exclude_hidden:
            dirs[:] = [d for d in dirs if not self._is_hidden_file(os.path.join(root, d))]
    def _generate_directory_entry(self, root, folder_path):
        """Generate tree entry for a directory"""
        level = root.replace(folder_path, '').count(os.sep)
        indent = ' ' * 4 * level
        return f'{indent}{os.path.basename(root)}/\n'

    def _generate_file_entries(self, files, root, folder_path):
        """Generate tree entries for files in a directory"""
        tree = ""
        level = root.replace(folder_path, '').count(os.sep)
        subindent = ' ' * 4 * (level + 1)
        for f in files:
            file_path = os.path.join(root, f)
            if not self._should_exclude(file_path, folder_path):
                tree += f'{subindent}{f}\n'
            elif self.verbose:
                print(f"Excluding file from tree: {file_path}")
                self.excluded_files_count += 1
        return tree

    def _log_tree_results(self, tree):
        """Log tree generation results if verbose mode is enabled"""
        if self.verbose:
            print(f"The file tree to be processed:\n{tree}")
            print(f"Total excluded items: {self.excluded_files_count}")

    def _get_file_contents(self, file_path):
        """Read file contents with better error handling"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            try:
                # Fall back to latin-1 for binary-like files
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    return f"[Binary/Non-UTF8 file - showing first 500 chars]\n{content[:500]}..."
            except OSError as e:
                return f"[Could not read file content: {str(e)}]"
        except OSError as e:
            return f"[Error reading file: {str(e)}]"

    def _is_hidden_file(self, file_path):
        """Check if file/directory is hidden"""
        components = os.path.normpath(file_path).split(os.sep)
        for c in components:
            if c.startswith((".", "__")):
                return True
        return False

    def _skip_excluded_dir(self, root, excluded_dirs):
        """Return True if the directory is inside an excluded path."""
        for excluded_dir in excluded_dirs:
            if root.startswith(excluded_dir):
                return True
        return False

    def _process_files(self, path):
        """Process files, respecting exclusion patterns"""
        content = ""
        processed_count = 0
        excluded_dirs = set()

        for root, dirs, files in os.walk(path):
            if self._handle_directory_exclusion(root, path, excluded_dirs):
                continue

            if self._skip_excluded_dir(root, excluded_dirs):
                continue

            self._filter_directories_for_processing(dirs, root, path)

            for file in files:
                file_result = self._process_single_file(file, root, path)
                if file_result:
                    content += file_result
                    processed_count += 1

        if self.verbose:
            print(f"Processed {processed_count} files")

        return content

    def _handle_directory_exclusion(self, root, path, excluded_dirs):
        """Handle directory exclusion for file processing"""
        if self._should_exclude(root, path):
            if self.verbose:
                print(f"Skipping excluded directory: {root}")
            excluded_dirs.add(root)
            return True
        return False

    def _filter_directories_for_processing(self, dirs, root, path):
        """Filter directories to avoid processing excluded ones"""
        original_dirs = dirs[:]
        dirs[:] = []
        for d in original_dirs:
            dir_path = os.path.join(root, d)
            if not self._should_exclude(dir_path, path):
                dirs.append(d)

    def _is_image_file(self, file_path):
        """Check if the file is an image based on extension"""
        return os.path.splitext(file_path)[1].lower() in self.image_extensions

    def _process_single_file(self, file, root, path):
        """Process a single file and return its content or None if excluded"""
        file_path = os.path.join(root, file)

        if self._should_exclude(file_path, path):
            if self.verbose:
                print(f"Skipping excluded file: {file_path}")
            return None

        if self.verbose:
            print(f"Processing: {file_path}")

        try:
            if self.output_type == 'docx' and self._is_image_file(file_path):
                # For images in docx mode, return special marker with evaluated path
                return f"\n\n(IMAGE_MARKER){os.path.abspath(file_path)}(/IMAGE_MARKER)\n"
            return self._format_file_content(file_path, path)
        except (OSError, UnicodeDecodeError) as e:
            return self._format_file_error(file_path, path, e)

    def _format_file_content(self, file_path, path):
        """Format successful file content for output"""
        file_content = self._get_file_contents(file_path)
        rel_path = os.path.relpath(file_path, path)

        content = f"\n\n{rel_path}\n"
        content += f"File type: {os.path.splitext(file_path)[1] or 'no extension'}\n"
        content += f"{file_content}"
        content += f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
        return content

    def _format_file_error(self, file_path, path, error):
        """Format file processing error for output"""
        if self.verbose:
            print(f"Error processing {file_path}: {error}")

        rel_path = os.path.relpath(file_path, path)
        content = f"\n\n{rel_path}\n"
        content += f"File type: {os.path.splitext(file_path)[1] or 'no extension'}\n"
        content += f"[Error: Could not process file - {str(error)}]"
        content += f"\n\n{'-' * 50}\nFile End\n{'-' * 50}\n"
        return content

    def get_text(self):
        """Generate the combined text output"""
        folder_structure = ""
        file_contents = ""

        if self.is_github_repo():
            self._clone_github_repo()
            folder_structure = self._parse_folder(self.temp_folder_path)
            file_contents = self._process_files(self.temp_folder_path)
        else:
            folder_structure = self._parse_folder(self.input_path)
            file_contents = self._process_files(self.input_path)

        # Section headers
        folder_structure_header = "Folder Structure"
        file_contents_header = "File Contents"

        # Delimiters
        delimiter = "-" * 50

        # Format the final text
        final_text = (f"{folder_structure_header}\n{delimiter}\n{folder_structure}\n\n"
                      f"{file_contents_header}\n{delimiter}\n{file_contents}")

        return final_text

    def get_file(self):
        """Generate and save the output file"""
        text = self.get_text()

        if self.output_type == "txt":
            with open(self.output_path, "w", encoding="utf-8") as file:
                file.write(text)
        elif self.output_type == "docx":
            doc = Document()
            # Split text into segments based on image markers
            segments = text.split('(IMAGE_MARKER)')
            
            # Add first text segment
            if segments[0].strip():
                doc.add_paragraph(segments[0])
            
            # Process remaining segments
            for segment in segments[1:]:
                if '(/IMAGE_MARKER)' in segment:
                    # Extract image path and remaining text
                    img_path, text_content = segment.split('(/IMAGE_MARKER)', 1)
                    img_path = str(img_path.strip())  # Remove any whitespace
                    if self.verbose:
                        print(f"Loading image from: {img_path}")
                    if not os.path.exists(img_path):
                        if self.verbose:
                            print(f"Image file not found at: {img_path}")
                        doc.add_paragraph(f"[Missing image: {img_path}]")
                    else:
                        try:
                            doc.add_picture(img_path, width=Inches(6))
                        except Exception as e:
                            if self.verbose:
                                print(f"Error adding image {img_path}: {e}")
                            doc.add_paragraph(f"[Error: Could not add image - {str(e)}]")
                    
                    if text_content.strip():
                        doc.add_paragraph(text_content)
            doc.save(self.output_path)
        else:
            raise ValueError("Invalid output type. Supported types: txt, docx")

        if self.verbose:
            print(f"Output saved to: {self.output_path}")

    #### GitHub Support ####

    def _clone_github_repo(self):
        """Clone GitHub repository to temporary directory"""
        try:
            self.temp_folder_path = tempfile.mkdtemp(prefix="github_repo_")
            git.Repo.clone_from(self.input_path, self.temp_folder_path)
            if self.verbose:
                print(f"GitHub repository cloned to: {self.temp_folder_path}")
        except Exception as e:
            print(f"Error cloning GitHub repository: {e}")
            raise

    def is_github_repo(self):
        """Check if input path is a GitHub repository URL"""
        return (self.input_path.startswith("https://github.com/") or
                self.input_path.startswith("git@github.com:"))

    def is_temp_folder_used(self):
        """Check if temporary folder is being used"""
        return self.temp_folder_path is not None

    def clean_up_temp_folder(self):
        """Clean up temporary folder"""
        if self.temp_folder_path and os.path.exists(self.temp_folder_path):
            shutil.rmtree(self.temp_folder_path)
            if self.verbose:
                print(f"Cleaned up temporary folder: {self.temp_folder_path}")


def main():
    """Main CLI entry point"""
    parser = argparse.ArgumentParser(
        description="Convert codebase to text with exclusion support.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s --input ./my_project --output output.txt --output_type txt
  %(prog)s --input https://github.com/user/repo --output repo.docx --output_type docx
  %(prog)s --input ./project --output out.txt --output_type txt --exclude "*.log,temp/,**/__pycache__/**"
  %(prog)s --input ./project --output out.txt --output_type txt --exclude "*.pyc" --exclude "build/"
        """
    )

    parser.add_argument("--input", help="Input path (folder or GitHub URL)", required=True)
    parser.add_argument("--output", help="Output file path", required=True)
    parser.add_argument("--output_type", help="Output file type (txt or docx)",
                       choices=["txt", "docx"],default="txt")
    parser.add_argument("--exclude", help="Exclude patterns (can be used multiple times)",
                       action="append", default=[])
    parser.add_argument("--exclude_hidden", help="Exclude hidden files and folders",
                       action="store_true")
    parser.add_argument("--verbose", help="Show detailed processing information",
                       action="store_true")

    args = parser.parse_args()

    try:
        code_to_text = CodebaseToText(
            input_path=args.input,
            output_path=args.output,
            output_type=args.output_type,
            verbose=args.verbose,
            exclude_hidden=args.exclude_hidden,
            exclude=args.exclude
        )

        code_to_text.get_file()

        if args.verbose:
            print("✅ Conversion completed successfully!")

    except (OSError, ValueError, git.GitCommandError) as e:
        print(f"❌ Error: {e}")
        return 1

    finally:
        # Clean up temporary folder if it was used
        if 'code_to_text' in locals() and code_to_text.is_temp_folder_used():
            code_to_text.clean_up_temp_folder()

    return 0


if __name__ == "__main__":
    sys.exit(main())

